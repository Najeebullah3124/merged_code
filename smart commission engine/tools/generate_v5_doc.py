from pathlib import Path
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "smart comimission engine doc V5.docx"

INCLUDE_EXTENSIONS = {
    ".js",
    ".json",
    ".md",
    ".py",
    ".yml",
    ".yaml",
    ".env",
}

EXCLUDE_PARTS = {
    "node_modules",
    ".git",
    "terminals",
    "agent-transcripts",
    "__pycache__",
}

PREFERRED_PREFIXES = [
    "README.md",
    ".env.example",
    "package.json",
    "Dockerfile",
    ".github/workflows/",
    "src/",
    "python/",
    "tools/",
    "infra/",
    "ops/",
    "docs/",
]


def should_include(path: Path) -> bool:
    rel = path.relative_to(ROOT)
    if any(part in EXCLUDE_PARTS for part in rel.parts):
        return False
    if path.name == "generate_v5_doc.py":
        return False
    if path.name == "package-lock.json":
        return False
    if path.name == "Dockerfile":
        return True
    return path.suffix in INCLUDE_EXTENSIONS


def explanation_for(rel_path: str) -> str:
    if rel_path.startswith("src/routes/"):
        return "Defines HTTP API endpoints and request/response orchestration."
    if rel_path.startswith("src/middleware/"):
        return "Implements cross-cutting request protections, security, and observability behavior."
    if rel_path.startswith("src/services/"):
        return "Contains reusable domain and infrastructure logic used by routes and runtime flows."
    if rel_path.startswith("src/models/"):
        return "Defines data schemas and persistence structures for MongoDB collections."
    if rel_path.startswith("src/config/"):
        return "Holds configurable runtime and policy defaults used across the service."
    if rel_path.startswith("src/db/"):
        return "Handles database connection and seeding/bootstrap responsibilities."
    if rel_path.startswith("src/state/"):
        return "Manages demo/offline state used when MongoDB is unavailable."
    if rel_path.startswith("python/"):
        return "Supports machine-learning and fraud-related simulation/microservice components."
    if rel_path.startswith("tools/"):
        return "Provides operational validation utilities such as load and chaos test scripts."
    if rel_path.startswith(".github/workflows/"):
        return "Defines CI/CD workflows for automated quality gates and deployment processes."
    if rel_path.startswith("infra/"):
        return "Infrastructure-as-code manifests for multi-environment deployment configuration."
    if rel_path.startswith("ops/alerts/"):
        return "Contains SRE alerting rules for production incident detection."
    if rel_path.startswith("ops/dashboards/"):
        return "Contains observability dashboard configuration for operations visibility."
    if rel_path.startswith("docs/"):
        return "Project documentation describing architecture, API behavior, and production evidence."
    if rel_path == "README.md":
        return "Primary project guide with setup, configuration, and operational references."
    if rel_path == "package.json":
        return "Node runtime manifest with scripts and dependency definitions."
    if rel_path == ".env.example":
        return "Template environment variables for local, staging, and production configuration."
    if rel_path == "Dockerfile":
        return "Container build specification for reproducible production runtime packaging."
    return "Source artifact included as part of the full project codebase."


def sort_key(rel_path: str):
    for i, prefix in enumerate(PREFERRED_PREFIXES):
        if rel_path.startswith(prefix):
            return (i, rel_path)
    return (len(PREFERRED_PREFIXES), rel_path)


def safe_read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def main():
    files = [p for p in ROOT.rglob("*") if p.is_file() and should_include(p)]
    rel_files = sorted([str(p.relative_to(ROOT)).replace("\\", "/") for p in files], key=sort_key)

    doc = Document()
    doc.add_heading("Smart Commission Engine - Full Code and Explanation", level=0)
    doc.add_paragraph("Document Version: V5")
    doc.add_paragraph(
        "This document contains the project code (source and operational artifacts) with a concise explanation for each file."
    )
    doc.add_paragraph(f"Total files included: {len(rel_files)}")

    for rel in rel_files:
        path = ROOT / rel
        doc.add_page_break()
        doc.add_heading(rel, level=1)
        doc.add_paragraph(f"Explanation: {explanation_for(rel)}")
        content = safe_read_text(path)
        doc.add_heading("Code", level=2)
        # Preserve plain text code blocks in Word.
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(content if content else "(File is empty)")
        run.font.name = "Consolas"

    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
