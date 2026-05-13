"""Convert BACKEND_API_REFERENCE.md to Word (.docx). Requires: pip install python-docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "BACKEND_API_REFERENCE.md"
# ASCII filename avoids shell / ZIP encoding issues on Windows.
OUT_PATH = ROOT / "generated-docx" / "Backend_API_Reference_one-platform.docx"


def strip_md_inline(s: str) -> str:
    """Remove common markdown emphasis markers for plain-text fallback."""
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


def add_paragraph_with_bold(
    doc: Document, text: str, style: str | None = None
) -> None:
    """Paragraph with **bold** segments as bold runs."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)


def parse_table_lines(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.startswith("|"):
            break
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    header = rows[0]
    # Skip separator row if present (|---|---|)
    data_start = 1
    if rows[1] and all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in rows[1]):
        data_start = 2
    data_rows = rows[data_start:]
    if not data_rows:
        return
    n_cols = max(len(header), max(len(r) for r in data_rows))
    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(header[:n_cols]):
        hdr_cells[j].text = strip_md_inline(h)
    for ri, row in enumerate(data_rows):
        for j in range(n_cols):
            val = row[j] if j < len(row) else ""
            table.rows[ri + 1].cells[j].text = strip_md_inline(val)


def convert(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if not line:
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("##"):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and "|" in lines[i + 1]:
            rows, end_i = parse_table_lines(lines, i)
            if len(rows) >= 2:
                add_markdown_table(doc, rows)
                i = end_i
                continue

        if line.startswith("- "):
            add_paragraph_with_bold(doc, line[2:], style="List Bullet")
            i += 1
            continue

        if (
            line.startswith("*")
            and line.endswith("*")
            and not line.startswith("**")
            and line.count("*") == 2
        ):
            p = doc.add_paragraph()
            r = p.add_run(strip_md_inline(line[1:-1]))
            r.italic = True
            i += 1
            continue

        add_paragraph_with_bold(doc, line)
        i += 1


def main() -> None:
    if not MD_PATH.is_file():
        raise SystemExit(f"Missing {MD_PATH}")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    convert(MD_PATH.read_text(encoding="utf-8"), doc)

    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
