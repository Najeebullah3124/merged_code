"""
Single Word export for the entire Simulation Intelligence service folder.

Produces ONE file: generated-docx/Simulation_Intelligence_Complete.docx

Contents:
  Part 1 — Full CTO platform design (docs/CTO_PLATFORM_DESIGN_SIMULATION_INTELLIGENCE.md)
  Part 2 — Full service reference markdown (docs/SIMULATION_INTELLIGENCE_COMPLETE_REFERENCE.md)
  Part 3 — Every repository file verbatim (app, public, tests, tools, docs, Docker, etc.)

Run:  python tools/build_documentation_docx.py
Requires: pip install python-docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "generated-docx" / "Simulation_Intelligence_Complete.docx"

SKIP_DIR_NAMES = frozenset(
    {"generated-docx", "__pycache__", ".git", ".pytest_cache", "node_modules", ".venv", "venv"}
)
SKIP_SUFFIXES = frozenset({".pyc", ".pyo", ".docx", ".tmp"})


def strip_md_inline(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


def add_paragraph_with_bold(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)


def parse_table_lines(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.startswith("|"):
            break
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    header = rows[0]
    data_start = 1
    if rows[1] and all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in rows[1]):
        data_start = 2
    data_rows = rows[data_start:]
    if not data_rows:
        return
    n_cols = max(len(header), max(len(r) for r in data_rows))
    table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    table.style = "Table Grid"
    for j, h in enumerate(header[:n_cols]):
        table.rows[0].cells[j].text = strip_md_inline(h)
    for ri, row in enumerate(data_rows):
        for j in range(n_cols):
            val = row[j] if j < len(row) else ""
            table.rows[ri + 1].cells[j].text = strip_md_inline(val)


def add_code_block_chunked(doc: Document, lines: list[str], font_size: int = 7) -> None:
    buf: list[str] = []
    for line in lines:
        buf.append(line)
        if len(buf) >= 45:
            p = doc.add_paragraph()
            r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(font_size)
            buf = []
    if buf:
        p = doc.add_paragraph()
        r = p.add_run("\n".join(buf))
        r.font.name = "Consolas"
        r.font.size = Pt(font_size)


def convert_markdown(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_acc: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if in_code:
            if line.strip().startswith("```"):
                in_code = False
                add_code_block_chunked(doc, code_acc)
                code_acc = []
                i += 1
                continue
            code_acc.append(raw)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = True
            code_acc = []
            i += 1
            continue

        if not line:
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("##"):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        if line.startswith("## ") and not line.startswith("###"):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue

        if line.startswith("### ") and not line.startswith("####"):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and "|" in lines[i + 1]:
            rows, end_i = parse_table_lines(lines, i)
            if len(rows) >= 2:
                add_markdown_table(doc, rows)
                i = end_i
                continue

        if re.match(r"^\d+\.\s", line):
            add_paragraph_with_bold(doc, re.sub(r"^\d+\.\s*", "", line), style="List Number")
            i += 1
            continue

        if line.startswith("- "):
            add_paragraph_with_bold(doc, line[2:], style="List Bullet")
            i += 1
            continue

        add_paragraph_with_bold(doc, line)
        i += 1

    if in_code and code_acc:
        add_code_block_chunked(doc, code_acc)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def should_include_file(path: Path) -> bool:
    try:
        rel = path.relative_to(ROOT)
    except ValueError:
        return False
    for part in rel.parts:
        if part in SKIP_DIR_NAMES:
            return False
    if path.suffix.lower() in SKIP_SUFFIXES:
        return False
    if path.name.startswith("."):
        if path.name in {".env.example", ".gitignore", ".dockerignore"}:
            return True
        return False
    return True


def iter_all_files() -> list[Path]:
    files: set[Path] = set()
    for p in ROOT.rglob("*"):
        if not p.is_file():
            continue
        if not should_include_file(p):
            continue
        if p.resolve() == OUT_PATH.resolve():
            continue
        files.add(p.resolve())
    for name in (".env.example", ".gitignore", ".dockerignore"):
        fp = (ROOT / name).resolve()
        if fp.is_file() and should_include_file(fp):
            files.add(fp)
    return sorted(files, key=lambda x: str(x.relative_to(ROOT)).replace("\\", "/"))


def read_file_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def add_verbatim_file(doc: Document, path: Path) -> None:
    rel = path.relative_to(ROOT)
    doc.add_heading(str(rel).replace("\\", "/"), level=2)
    meta = doc.add_paragraph()
    meta.add_run("Full path: ").bold = True
    meta.add_run(str(path))
    try:
        body = read_file_text(path)
    except OSError as e:
        doc.add_paragraph(f"[Read error: {e}]", style="Intense Quote")
        return
    lines = body.splitlines()
    if not lines:
        doc.add_paragraph("(empty file)")
        return
    add_code_block_chunked(doc, lines, font_size=7)


def build() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        print("pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Simulation Intelligence Platform — Complete Documentation")
    r.bold = True
    r.font.size = Pt(22)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(
        "Single export: CTO design + technical reference + full repository source"
    ).italic = True
    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.add_run("Document structure:").bold = True
    doc.add_paragraph(
        "Part 1 — Staff/CTO platform design. Part 2 — Service reference (markdown). "
        "Part 3 — Verbatim copy of every project file.",
        style="List Bullet",
    )

    p1 = ROOT / "docs" / "CTO_PLATFORM_DESIGN_SIMULATION_INTELLIGENCE.md"
    doc.add_heading("Part 1 — CTO platform design (complete)", level=0)
    if p1.is_file():
        convert_markdown(doc, read_file_text(p1))
    else:
        doc.add_paragraph(f"(Missing: {p1})")

    add_page_break(doc)

    p2 = ROOT / "docs" / "SIMULATION_INTELLIGENCE_COMPLETE_REFERENCE.md"
    doc.add_heading("Part 2 — Simulation Intelligence Service reference (complete)", level=0)
    if p2.is_file():
        convert_markdown(doc, read_file_text(p2))
    else:
        doc.add_paragraph(f"(Missing: {p2})")

    add_page_break(doc)

    doc.add_heading("Part 3 — Repository files (verbatim, complete)", level=0)
    doc.add_paragraph(
        "All files under this service (excluding build output, caches, and VCS). "
        "Build this document with: python tools/build_documentation_docx.py"
    )

    all_paths = iter_all_files()
    doc.add_paragraph(f"File count: {len(all_paths)}", style="Intense Quote")

    for path in all_paths:
        add_verbatim_file(doc, path)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")
    print(f"Part 3 file count: {len(all_paths)}")


if __name__ == "__main__":
    build()
