r"""
Build one complete Word document for all smart-pricing code.

Output:
    D:\New Forecasting\Merge Code\smart pricing V6.docx

What it includes:
    - Every text/code file under smart pricing projects (verbatim; no truncation)
    - A short explanation section for each file
    - Project-level overview and file inventory

Run:
    python build_smart_pricing_v6_docx.py
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(r"D:\New Forecasting\Merge Code")
OUT_PATH = ROOT / "smart pricing V6.docx"

SKIP_DIRS = {
    ".git",
    "__pycache__",
    ".pytest_cache",
    "node_modules",
    ".next",
    ".venv",
    "venv",
    "dist",
    "build",
    "generated-docx",
}

TEXT_SUFFIXES = {
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".json",
    ".md",
    ".txt",
    ".yml",
    ".yaml",
    ".toml",
    ".ini",
    ".cfg",
    ".env",
    ".html",
    ".css",
    ".sql",
    ".sh",
    ".bat",
    ".ps1",
    ".xml",
    ".csv",
    ".gitignore",
    ".dockerignore",
}

TEXT_FILENAMES = {
    "Dockerfile",
    "dockerfile",
    "Makefile",
    "makefile",
    "requirements",
    "requirements.txt",
    "package.json",
    "package-lock.json",
    "pnpm-lock.yaml",
    "yarn.lock",
}


@dataclass
class FileEntry:
    path: Path
    rel: str
    line_count: int


def find_smart_pricing_roots(root: Path) -> list[Path]:
    roots: list[Path] = []
    for p in root.iterdir():
        if not p.is_dir():
            continue
        name = p.name.lower()
        if "smart" in name and ("pricing" in name or "prcing" in name):
            roots.append(p)
    return sorted(roots, key=lambda x: x.name.lower())


def is_text_code_file(path: Path) -> bool:
    if path.name in TEXT_FILENAMES:
        return True
    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return True
    return False


def iter_code_files(project_root: Path) -> Iterable[Path]:
    for p in project_root.rglob("*"):
        if not p.is_file():
            continue
        rel_parts = p.relative_to(project_root).parts
        if any(part in SKIP_DIRS for part in rel_parts):
            continue
        if not is_text_code_file(p):
            continue
        yield p


def safe_read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True


def add_code_block(doc: Document, content: str) -> None:
    lines = content.splitlines()
    if not lines:
        doc.add_paragraph("(empty file)")
        return
    chunk: list[str] = []
    for line in lines:
        chunk.append(line)
        if len(chunk) >= 45:
            para = doc.add_paragraph()
            run = para.add_run("\n".join(chunk))
            run.font.name = "Consolas"
            run.font.size = Pt(7)
            chunk = []
    if chunk:
        para = doc.add_paragraph()
        run = para.add_run("\n".join(chunk))
        run.font.name = "Consolas"
        run.font.size = Pt(7)


def explain_file(rel_path: str) -> str:
    low = rel_path.lower()
    if "/api" in low or "main.py" in low or "server" in low:
        return "Entry-point/API layer file. It defines service routes, request handling, and runtime wiring."
    if "engine" in low or "optimizer" in low or "decision" in low:
        return "Core decision/simulation logic file. It contains business and optimization behavior."
    if "model" in low or "feature" in low or "inference" in low:
        return "ML/data modeling file. It covers feature handling, model interaction, or prediction utilities."
    if "test" in low:
        return "Test coverage file. It validates behavior and protects against regressions."
    if low.endswith(".md"):
        return "Documentation file included verbatim for full technical context."
    if low.endswith(".json") or "config" in low:
        return "Configuration/schema file used by runtime, tooling, or deployment."
    if low.endswith(".yml") or low.endswith(".yaml") or "docker" in low:
        return "Deployment/infra configuration file (containers, orchestration, CI/CD, or ops)."
    return "Implementation/support file included verbatim."


def build() -> None:
    projects = find_smart_pricing_roots(ROOT)
    if not projects:
        raise SystemExit("No smart pricing directories found.")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    add_title(doc, "Smart Pricing V6")
    add_subtitle(
        doc,
        "Complete code compendium with explanation — every included file is verbatim, line-for-line.",
    )
    doc.add_paragraph()

    doc.add_heading("Included smart pricing projects", level=1)
    for pr in projects:
        doc.add_paragraph(f"- {pr.name}", style="List Bullet")

    all_entries: list[FileEntry] = []

    for project in projects:
        project_files = sorted(iter_code_files(project), key=lambda x: str(x.relative_to(ROOT)).lower())
        doc.add_heading(f"Project: {project.name}", level=1)
        doc.add_paragraph(
            "This section includes all detected text/code files under the project folder. "
            "Each file appears with a short purpose statement and full raw content."
        )
        doc.add_paragraph(f"File count: {len(project_files)}", style="Intense Quote")

        for file_path in project_files:
            rel = str(file_path.relative_to(ROOT)).replace("\\", "/")
            content = safe_read_text(file_path)
            line_count = len(content.splitlines())
            all_entries.append(FileEntry(path=file_path, rel=rel, line_count=line_count))

            doc.add_heading(rel, level=2)
            doc.add_paragraph(explain_file(rel), style="Intense Quote")
            doc.add_paragraph(f"Lines: {line_count}")
            add_code_block(doc, content)

    doc.add_heading("Global inventory", level=1)
    doc.add_paragraph(
        f"Total projects: {len(projects)} | Total files: {len(all_entries)} | "
        f"Total lines: {sum(e.line_count for e in all_entries)}"
    )
    for entry in all_entries:
        doc.add_paragraph(f"- {entry.rel} ({entry.line_count} lines)", style="List Bullet")

    doc.save(str(OUT_PATH))
    print(f"Wrote {OUT_PATH}")
    print(f"Projects: {len(projects)} | Files: {len(all_entries)} | Lines: {sum(e.line_count for e in all_entries)}")


if __name__ == "__main__":
    build()

