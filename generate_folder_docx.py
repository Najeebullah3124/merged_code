from __future__ import annotations

from pathlib import Path
import re

from docx import Document


ROOT = Path(r"D:/New Forecasting/Merge Code")
EXCLUDED_TOP_LEVEL = {"all merge code", "one-platform", ".git", ".vscode", "__pycache__"}
OUTPUT_DIR = ROOT / "generated-docx"

# Common code/text-like extensions to include as "full code"
TEXT_EXTENSIONS = {
    ".py",
    ".ts",
    ".tsx",
    ".js",
    ".jsx",
    ".java",
    ".go",
    ".rs",
    ".rb",
    ".php",
    ".cs",
    ".cpp",
    ".c",
    ".h",
    ".hpp",
    ".sql",
    ".sh",
    ".ps1",
    ".yml",
    ".yaml",
    ".json",
    ".toml",
    ".ini",
    ".cfg",
    ".conf",
    ".md",
    ".txt",
    ".xml",
    ".html",
    ".css",
    ".scss",
    ".env",
    ".dockerfile",
}

SPECIAL_FILENAMES = {
    "Dockerfile",
    ".dockerignore",
    ".gitignore",
    ".env",
    ".env.example",
    "Makefile",
    "requirements.txt",
    "package.json",
    "package-lock.json",
    "pnpm-lock.yaml",
    "yarn.lock",
    "README.md",
}

SKIP_DIR_NAMES = {
    ".git",
    ".idea",
    ".vscode",
    "__pycache__",
    "node_modules",
    "dist",
    "build",
    ".next",
    ".turbo",
    ".venv",
    "venv",
    ".pytest_cache",
    ".mypy_cache",
    "coverage",
    "generated-docx",
}


def sanitize_filename(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*]+', "_", name).strip(" .")


def is_probably_text(file_path: Path) -> bool:
    name = file_path.name
    suffix = file_path.suffix.lower()
    if name in SPECIAL_FILENAMES:
        return True
    if suffix in TEXT_EXTENSIONS:
        return True
    return False


def safe_read_text(file_path: Path) -> str | None:
    try:
        raw = file_path.read_bytes()
    except Exception:
        return None

    # Skip very large files to keep docx generation stable
    if len(raw) > 1_500_000:
        return f"[Skipped content: file too large ({len(raw)} bytes)]"

    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except Exception:
            continue
    return None


def folder_explanation(folder_name: str, file_count: int) -> str:
    return (
        f"This document covers the folder '{folder_name}'. "
        f"It contains an overview and full code listings for {file_count} text/code files "
        "found in this folder (excluding heavy generated/dependency directories)."
    )


def collect_text_files(folder: Path) -> list[Path]:
    files: list[Path] = []
    for path in folder.rglob("*"):
        if path.is_dir():
            continue
        parts_lower = {part.lower() for part in path.parts}
        if parts_lower & {d.lower() for d in SKIP_DIR_NAMES}:
            continue
        if is_probably_text(path):
            files.append(path)
    return sorted(files)


def create_doc_for_folder(folder: Path) -> Path:
    files = collect_text_files(folder)
    doc = Document()
    doc.add_heading(f"Codebook - {folder.name}", level=1)
    doc.add_heading("Explanation", level=2)
    doc.add_paragraph(folder_explanation(folder.name, len(files)))
    doc.add_paragraph(
        "Note: Binary files and large generated/dependency directories are excluded to keep "
        "the document readable and stable."
    )

    doc.add_heading("Full Code", level=2)
    if not files:
        doc.add_paragraph("No eligible text/code files found.")
    for file_path in files:
        rel = file_path.relative_to(folder)
        doc.add_heading(str(rel), level=3)
        content = safe_read_text(file_path)
        if content is None:
            doc.add_paragraph("[Skipped: could not decode as text]")
        else:
            doc.add_paragraph(content)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"{sanitize_filename(folder.name)}.docx"
    doc.save(output_path)
    return output_path


def main() -> None:
    top_dirs = [p for p in ROOT.iterdir() if p.is_dir() and p.name not in EXCLUDED_TOP_LEVEL]
    for folder in sorted(top_dirs, key=lambda p: p.name.lower()):
        output = create_doc_for_folder(folder)
        print(f"Created: {output}")


if __name__ == "__main__":
    main()
